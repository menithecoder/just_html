import fitz  # PyMuPDF for PDF handling
from PIL import Image
from docx import Document
import os
from pyppeteer import launch
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
def pdf_to_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    images = []

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        image = page.get_pixmap()
        img = Image.frombytes("RGB", [image.width, image.height], image.samples)
        images.append(img)

    pdf_document.close()
    return images


def save_images(images, output_format='PNG'):
    download_path = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Ensure the directory exists, create it if necessary
    if not os.path.exists(download_path):
        os.makedirs(download_path)

    for i, img in enumerate(images):
        img.save(os.path.join(download_path, f"output_{i}.{output_format.lower()}"), format=output_format)



def pdf_to_text(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = ''
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    pdf_document.close()
    print(text)
    return text

def text_to_word(text, output_path='output.docx'):
    download_path = os.path.join(os.path.expanduser("~"), "Downloads", output_path)
    doc = Document()
    doc.add_paragraph(text)
    doc.save(download_path)
###################################
#word to other
####################################
def convert_docx_to_pdf(input_docx, output_pdf):
    doc = Document(input_docx)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
    </head>
    <body>
        {text}
    </body>
    </html>
    '''

    async def generate_pdf():
        browser = await launch()
        page = await browser.newPage()
        await page.setContent(html)
        await page.pdf({'path': output_pdf, 'format': 'A4'})
        await browser.close()

    import asyncio
    asyncio.get_event_loop().run_until_complete(generate_pdf())


def convert_docx_to_jpeg(input_docx, output_,format_type):
    # Read the DOCX file
    doc = Document(input_docx)

    # Create an image
    img_width = 800
    img_height = 600
    background_color = 'white'
    font_size = 20
    font_color = 'black'

    image = Image.new('RGB', (img_width, img_height), color=background_color)
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("arial.ttf", font_size)

    # Convert the text content from the DOCX to image
    text_content = '\n'.join([para.text for para in doc.paragraphs])
    draw.text((10, 10), text_content, font=font, fill=font_color)
    output_t = os.path.join(os.path.expanduser("~"), "Downloads", output_)
    # Save the image as JPEG
    image.save(output_t, format_type)

def convert_png_to_pdf(input_path, output_filename):
    image = Image.open(input_path)

    # Get the path to the user's download directory
    download_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Create the full output path including the download directory
    output_path = os.path.join(download_dir, output_filename)

    # Create a PDF
    pdf = canvas.Canvas(output_path, pagesize=image.size)

    # Draw the PNG image onto the PDF
    pdf.drawImage(input_path, 0, 0, width=image.width, height=image.height)

    # Save the PDF file
    pdf.save()
    print(f"Converted {input_path} to {output_path}")

def convert_png_to_jpeg(input_path, output_path):
    try:
        # Open the PNG image
        with Image.open(input_path) as img:
            # Convert PNG to JPEG
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Save the image as JPEG
            img.save(output_path, 'JPEG')

            print(f"File saved successfully as {output_path}")
    except Exception as e:
        print(f"Error converting file: {e}")

def convert_jpeg_to_png(input_path, output_path):
    try:
        # Open the JPEG image
        with Image.open(input_path) as img:
            # Convert JPEG to PNG
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Save the image as PNG
            img.save(output_path, 'PNG')

            print(f"File saved successfully as {output_path}")
    except Exception as e:
        print(f"Error converting file: {e}")
def convert_jpeg_to_pdf(input_image, output_pdf):
    # Create a new PDF file
    c = canvas.Canvas(output_pdf, pagesize=letter)

    # Open the image using PIL
    img = Image.open(input_image)

    # Get the dimensions of the image
    width, height = img.size

    # Add the image to the PDF file
    c.setPageSize((width, height))
    c.drawImage(input_image, 0, 0, width, height)
    c.showPage()

    # Save the PDF file
    c.save()



###################################
#pdf to other
####################################
# Example usage:
pdf_file_path = input("Enter format to convert  to")
format=input("Enter format to convert  to")

if pdf_file_path.lower().endswith('.pdf'):

    if format == "PNG" or format == "JPEG":
        # Convert PDF to images (PNG by default)
        pdf_images = pdf_to_images(pdf_file_path)
        if format == "PNG":
            # Save images in different formats (PNG, JPEG)
            save_images(pdf_images, output_format='PNG')
        elif format == "JPEG":
            save_images(pdf_images, output_format='JPEG')
    elif format=="DOCX":
        extracted_text = pdf_to_text(pdf_file_path)
        text_to_word(extracted_text)


###################################
#word to other
####################################




elif pdf_file_path.lower().endswith('.docx'):
    output_image_file = 'convert{}.png'

    if format=="PDF":
        output_pdf_file = 'converted.pdf'  # Replace with your desired output PDF filename
        download_path = os.path.join(os.path.expanduser("~"), "Downloads", output_pdf_file)
        convert_docx_to_pdf(pdf_file_path, download_path)

    elif format=="PNG":
        output_png_path = 'output_image.png'
        convert_docx_to_jpeg(pdf_file_path, output_png_path,'png')

    elif format=="JPEG":
        output_jpeg_path = 'output_image.jpg'
        convert_docx_to_jpeg(pdf_file_path, output_jpeg_path,'jpeg')
###################################
# png to other
####################################
elif pdf_file_path.lower().endswith('.png'):
    if format=="PDF":
        output_file = 'output.pdf'
        convert_png_to_pdf(pdf_file_path, output_file)
    elif format=="JPEG":
        download_directory = os.path.expanduser('~') + '/Downloads/'
        output_file_path = os.path.join(download_directory, 'output_file.jpg')

        # Convert PNG to JPEG and save in the download directory
        convert_png_to_jpeg(pdf_file_path, output_file_path)



if pdf_file_path.lower().endswith(('.jpg', '.jpeg')):
    if format=="PDF":
        # Get the path to the Downloads directory
        download_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

        # Replace 'output_pdf_file.pdf' with the desired output PDF file name
        output_pdf_file = os.path.join(download_dir, 'output_pdf_file.pdf')

        # Convert the file extension to lowercase for comparison

        convert_jpeg_to_pdf(pdf_file_path, output_pdf_file)
    if format=="PNG":
        downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

        # Input JPEG file path


        # Output PNG file path in the Downloads directory
        output_png_path = os.path.join(downloads_dir, 'converted_image.png')

        # Call the function to convert JPEG to PNG
        convert_jpeg_to_png(pdf_file_path, output_png_path)


























